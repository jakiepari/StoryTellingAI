from langchain_ollama import OllamaLLM
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.prompts import PromptTemplate
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

def init_ollama():
    return OllamaLLM(
        model="llama3.2:1b-instruct-q5_K_S", 
        callbacks=[StreamingStdOutCallbackHandler()],
        temperature=0.9,
        num_ctx=2048,
        num_thread=2,
        timeout=60
    )

STORY_PROMPT = """
You are a master storyteller AI with deep knowledge of narrative structures. Create a captivating story based on the following structure. Each part should be unique and contribute to a well-rounded, engaging narrative.

Instructions:
- **Title:** A unique title representing the theme of the story.
- **Genre:** List the genres that best describe the story.
- **Introduction Characters:** List each main character with their name, age, role, and brief description.
- **Prologue:** Introduce the setting and main premise, giving readers a sense of the world and context.
- **Main Story:**
    Create chapters with at least 2 detailed paragraphs each (10-40 paragraphs total):
    - Each chapter must have a title
    - Each paragraph should be at least 5 sentences long
    - Paragraphs should focus on different aspects: action, dialogue, character development, setting descriptions, internal thoughts
    - Ensure smooth transitions between paragraphs and chapters
    - Build tension and develop conflicts throughout
    - Include meaningful character interactions and plot developments
    - Maintain consistent pacing and tone
    - Each chapter should advance the story in multiple ways
- **Conclusion:** Provide a satisfying ending that resolves the main conflicts

**Paragraph Structure Guidelines:**
- Each paragraph must be completely unique and advance the story in a new way
- Do not reuse or loop content from previous paragraphs
- Include vivid descriptions, emotional depth, and engaging dialogue
- Maintain continuity while introducing new elements
- Focus on showing rather than telling
- Vary sentence structure and length for better flow

**Note:** Focus solely on delivering the story following the format:
- Title
- Genre
- Introduction Characters  
- Prologue
- Main Story (Multiple chapters with detailed paragraphs)
- Conclusion

Additional Details:
- Story Type: {story_type}
- Input Type: {input_type}
- Input Value: {input_value}
- Additional Instructions: {additional_instructions}
"""

REVISION_PROMPT = """
Please revise and improve the story based on the following feedback:

Original Story:
{original_story}

User Feedback:
{feedback}

Please maintain the same structure but incorporate the requested changes and improvements while enhancing the overall narrative quality.
"""

class StoryTellerAI:
    def __init__(self):
        self.llm = init_ollama()
        self.genres = [
            "Fantasy", "Science Fiction", "Mystery", "Romance", 
            "Adventure", "Horror", "Historical Fiction", "Comedy",
            "Thriller", "Drama", "Dystopian", "Paranormal",
            "Young Adult", "Children's Fiction", "Crime", "Superhero",
            "Western", "Slice of Life", "Fairy Tale", "Mythology",
            "Urban Fantasy", "Historical Romance", "Psychological Thriller",
            "Dark Fantasy", "Epic Fantasy", "Space Opera", "Post-Apocalyptic", 
            "Steampunk", "Cyberpunk", "Detective Fiction", "Medical Fiction",
            "Political Fiction", "Satire", "Military Fiction", "Gothic",
            "Spy Fiction", "Legal Thriller", "Biographical Fiction",
            "Coming of Age", "Family Saga", "Time Travel", "Road Trip"
        ]
        # Track generated content to avoid repetition
        self.used_content = set()
        self.selected_genres = []

    def get_user_preferences(self):
        concept = input("Enter a story concept (press Enter to skip): ").strip()
        title = "" if concept else input("Enter a story title (press Enter for genre selection): ").strip()
        
        if not concept and not title:
            print("\nAvailable genres (select max is 5):")
            for i, genre in enumerate(self.genres, 1):
                print(f"{i}. {genre}")
            
            while len(self.selected_genres) < 5:
                try:
                    genre_idx = int(input(f"\nSelect genre number (1-{len(self.genres)}, or 0 to finish): ")) 
                    if genre_idx == 0:
                        break
                    if 1 <= genre_idx <= len(self.genres):
                        selected_genre = self.genres[genre_idx-1]
                        if selected_genre not in self.selected_genres:
                            self.selected_genres.append(selected_genre)
                            print(f"Selected genres: {', '.join(self.selected_genres)}")
                        else:
                            print("Genre already selected. Please choose another.")
                    else:
                        print("Invalid genre number. Please try again.")
                except ValueError:
                    print("Please enter a valid number.")
            
            genres_str = " and ".join(self.selected_genres)
            concept = f"A story combining elements of {genres_str}"
        
        return {
            "concept": concept,
            "title": title
        }
    def save_to_word(self, story, filename):
        doc = Document()
        doc.add_heading('AI Generated Story', 0)
        doc.add_paragraph(story)
        doc.save(f"{filename}.docx")
        print(f"\nStory saved to {filename}.docx")

    def save_to_pdf(self, story, filename):
        c = canvas.Canvas(f"{filename}.pdf", pagesize=A4)
        
        c.setFont("Helvetica-Bold", 24)
        c.drawString(50, A4[1] - 50, "AI Generated Story")
        
        text = c.beginText()
        text.setTextOrigin(50, A4[1] - 100)
        text.setFont("Helvetica", 12)
        text.setLeading(14)
        
        lines = story.splitlines()
        for line in lines:
            words = line.split()
            current_line = []
            for word in words:
                current_line.append(word)
                if c.stringWidth(' '.join(current_line), "Helvetica", 12) > A4[0] - 100:
                    current_line.pop()
                    text.textLine(' '.join(current_line))
                    current_line = [word]
                    
            if current_line:
                text.textLine(' '.join(current_line))
            
            if text.getY() < 50:
                c.drawText(text)
                c.showPage()
                text = c.beginText()
                text.setTextOrigin(50, A4[1] - 50)
                text.setFont("Helvetica", 12)
                text.setLeading(14)
        
        c.drawText(text)
        c.save()
        print(f"\nStory saved to {filename}.pdf")

    def check_content_uniqueness(self, content):
        # Convert content to lowercase and remove punctuation for comparison
        normalized = ''.join(c.lower() for c in content if c.isalnum())
        content_parts = set(normalized[i:i+50] for i in range(len(normalized)-49))
        
        # Check for significant overlap with previously used content
        overlap = content_parts & self.used_content
        if len(overlap) / len(content_parts) > 0.3:  # More than 30% overlap
            return False
            
        # Add new content parts to used content
        self.used_content.update(content_parts)
        return True

    def generate_story(self, preferences):
        try:
            # Add anti-repetition instruction to prompt
            anti_repetition_instruction = """
            Important: Each section must be completely unique. Do not repeat plot points, 
            descriptions, or character developments. Follow a clear narrative progression 
            without loops or redundancy. Each paragraph should advance the story in a new direction.
            """
            
            prompt = PromptTemplate(
                template=STORY_PROMPT + anti_repetition_instruction,
                input_variables=["input_type", "input_value", "story_type", "additional_instructions"]
            )
            
            story_prompt = prompt.format(
                input_type="concept" if preferences["concept"] else "title",
                input_value=preferences["concept"] if preferences["concept"] else preferences["title"],
                story_type="story",
                additional_instructions="Create a coherent story with unique progression in each section"
            )
            
            # Generate story and check for uniqueness
            attempts = 0
            while attempts < 3:
                full_story = self.llm.invoke(story_prompt)
                
                if len(full_story.split()) < 50:
                    attempts += 1
                    continue
                    
                if self.check_content_uniqueness(full_story):
                    break
                    
                attempts += 1
                story_prompt += "\nPlease generate a completely different story with no repeated elements."
            
            print(full_story)

            while True:
                improve = input("\nWould you like to improve or fix anything in the story? (Enter feedback or 'skip' to continue): ").strip()
                if improve.lower() == 'skip' or not improve:
                    break
                    
                revision_prompt = PromptTemplate(
                    template=REVISION_PROMPT,
                    input_variables=["original_story", "feedback"]
                )
                
                revised_story = self.llm.invoke(revision_prompt.format(
                    original_story=full_story,
                    feedback=improve
                ))
                
                if self.check_content_uniqueness(revised_story):
                    full_story = revised_story
                    print("\nRevised story:\n")
                    print(full_story)
                else:
                    print("\nRevision contained too much repeated content. Keeping original version.")

            save = input("\nWould you like to save this story? (word/pdf/no): ").lower()
            if save == "word":
                filename = input("Enter filename (without extension): ")
                self.save_to_word(full_story, filename)
            elif save == "pdf":
                filename = input("Enter filename (without extension): ")
                self.save_to_pdf(full_story, filename)
            
            return full_story
            
        except Exception as e:
            print(f"Error generating story: {str(e)}")
            return "Failed to generate story. Please try again."

def main():
    storyteller = StoryTellerAI()
    print("Welcome to AI Story Generator!")
    print("Welcome ke AI Story Generator!")
    print("=" * 50)
    
    while True:
        try:
            preferences = storyteller.get_user_preferences()
            print("\nGenerating your story...")
            story = storyteller.generate_story(preferences)
            
            if input("\nGenerate another story? (y/n): ").lower() != 'y':
                break
                
        except Exception as e:
            print(f"An error occurred: {str(e)}")
            if input("\nTry again? (y/n): ").lower() != 'y':
                break
    
    print("\nThank you for using AI Story Generator!\nCredit:https://github.com/jakiepari/StoryTellingAI ")

if __name__ == "__main__":
    main()
